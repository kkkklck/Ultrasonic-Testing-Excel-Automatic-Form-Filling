# -*- coding: utf-8 -*-
# made by lck
# 功能说明：从Word探伤报告中提取关键信息，自动填充到Excel模板中，支持多日期数据拆分与批量填写
# 版本：v4.0.1


# 环境说明：(使用须知！！！！！！很重要，必看！！！！！）
# 1. 运行依赖：需安装 Python 3.6 及以上版本（推荐 3.8+，兼容性更优）
# 2. 必要库：需提前安装处理Word和Excel的专用库，安装命令：
#    pip install python-docx pywin32
#    若安装速度慢，可使用清华大学镜像：
#    pip install python-docx pywin32 -i https://pypi.tuna.tsinghua.edu.cn/simple
# 3. 系统兼容性：仅支持 Windows 系统（依赖 pywin32 操作Excel COM接口），文件路径需按Windows格式填写：
#    - 示例：D:\pycharm\存放docx专用\2025-046111\报告.docx、C:\Users\用户名\Desktop\数据.xlsx
# 4. 注意事项：
#    - Word 需为 .docx 格式，Excel 模板及数据文件支持 .xlsx/.xls 格式
#    - 运行前需确保：目标Excel模板已在Excel中打开并激活（程序会等待激活状态）
#    - 外部数据Excel的日期标记需在 F 列（如“3.31”“4/4”“4月4日”），用于拆分多日期数据
#    - 仪器编号自动判断：2025-03-12 至 2025-04-09 期间用“13-27”，其他时间用“22-72”
#    - 检测环境温度为北京月均温插值+微扰（纯数字），非实际测量值，仅供参考
#    - 多日期场景下，需手动在Excel中新建下一张模板并激活后再继续填写
#    - 路径输入支持“q”退出程序，文件不存在时会提示重试


# 导入必要库：上下文管理、正则、系统操作、时间、文件操作、日期处理、类型提示、路径处理、Excel/Word操作
from contextlib import contextmanager
import re
import sys
import time
import shutil
from datetime import datetime
from typing import Optional, List, Tuple, Iterable, Union
from re import Match as ReMatch
from pathlib import Path
import os
import win32com.client as win32  # 用于操作Excel COM接口（Windows专属）
from docx import Document  # 用于解析Word文档

# made by lck, an intern of this company in 2025 summer

# ------------------- 控制台心跳/进度条 -------------------
def print_step(msg: str) -> None:
    """打印带时间戳的步骤信息，用于跟踪程序执行进度"""
    now = time.strftime("%H:%M:%S")  # 格式化当前时间为时分秒
    print(f"[{now}] {msg}", flush=True)  # 强制刷新输出，确保实时显示


def progress_bar(done: int, total: int, prefix: str = "", bar_len: int | None = None) -> None:
    """
    显示进度条，直观展示任务完成比例
    :param done: 已完成数量
    :param total: 总数量
    :param prefix: 进度条前缀文本
    :param bar_len: 进度条长度，默认自适应终端宽度
    """
    if total <= 0:
        return
    # 自动计算进度条长度（适应终端宽度，限制在10-40字符）
    if bar_len is None:
        try:
            cols = shutil.get_terminal_size((80, 20)).columns  # 获取终端宽度
        except Exception:
            cols = 80
        bar_len = max(10, min(40, cols - len(prefix) - 20))  # 预留前缀和百分比的位置
    ratio = min(max(done / total, 0.0), 1.0)  # 计算完成比例（限制在0-1之间）
    filled = int(bar_len * ratio)  # 已填充的进度条长度
    bar = "█" * filled + "░" * (bar_len - filled)  # 进度条字符（█表示完成，░表示未完成）
    pct = int(ratio * 100)  # 百分比
    end = "\n" if done >= total else ""  # 完成时换行，否则不换行（覆盖当前行）
    print(f"\r{prefix} [{bar}] {pct:3d}% ({done}/{total})", end=end, flush=True)


# ------------------- 输入/等待工具 -------------------
@contextmanager
def excel_quiet(app):
    """
    上下文管理器：临时关闭Excel的事件响应、弹窗和屏幕刷新
    用途：加速Excel操作，避免插件弹窗干扰
    """
    # 保存原始设置
    old_alerts = app.DisplayAlerts
    old_events = app.EnableEvents
    old_update = app.ScreenUpdating
    try:
        # 关闭干扰项
        app.DisplayAlerts = False
        app.EnableEvents = False
        app.ScreenUpdating = False
        yield  # 执行with块内的代码
    finally:
        # 恢复原始设置
        app.ScreenUpdating = old_update
        app.EnableEvents = old_events
        app.DisplayAlerts = old_alerts


def get_desktop_path() -> Path:
    """
    获取系统桌面路径（兼容OneDrive桌面和普通桌面）
    优先级：OneDrive桌面 → 用户目录桌面 → 主目录桌面 → 主目录
    """
    candidates = [
        Path(os.path.expandvars(r"%OneDrive%")) / "Desktop",  # OneDrive桌面（若存在）
        Path(os.path.expandvars(r"%USERPROFILE%")) / "Desktop",  # 用户目录桌面
        Path.home() / "Desktop",  # 主目录桌面
    ]
    for p in candidates:
        try:
            if p.exists():  # 检查路径是否存在
                return p
        except Exception:
            pass
    return Path.home()  # 兜底：返回主目录


def prompt_path_with_retry(
    prompt_text: str,
    default_path: Optional[Path],
    must_exist: bool = True,
    allowed_suffixes: Optional[Iterable[str]] = None
) -> Path:  # noqa
    """
    循环提示用户输入文件路径，直到符合要求（支持默认路径、格式校验、存在性校验）
    :param prompt_text: 提示文本
    :param default_path: 默认路径（回车时使用）
    :param must_exist: 是否要求文件必须存在
    :param allowed_suffixes: 允许的文件后缀（如[".docx"]）
    :return: 符合要求的路径
    """
    allowed = {s.lower() for s in (allowed_suffixes or [])}  # 统一转为小写后缀
    while True:
        raw = input(f"{prompt_text}（回车默认：{default_path}）：\n→ ").strip().strip('"')
        if raw.lower() == "q":  # 输入q退出程序
            print("👋 已退出。")
            sys.exit(1)
        # 处理用户输入（为空则用默认路径）
        if raw:
            p: Path = Path(raw)
        else:
            if default_path is None:
                print("⚠️ 没有默认路径，请手动输入有效文件或按 q 退出。")
                continue
            p = default_path
        # 校验文件类型
        if allowed and p.suffix.lower() not in allowed:
            print(f"❌ 不支持的文件类型：{p.suffix}，仅支持：{', '.join(sorted(allowed))}")
            continue
        # 校验文件是否存在
        if must_exist and not p.exists():
            print(f"❌ 找不到文件：{p}，请重输（或按 q 退出）")
            continue
        return p


def wait_active_workbook(excel_app, tip: str):
    """
    等待用户在Excel中打开并激活目标工作簿（循环检查，直到有激活的工作簿）
    :param excel_app: Excel应用实例
    :param tip: 提示用户的文本
    :return: 激活的工作簿对象
    """
    while True:
        wb = excel_app.ActiveWorkbook  # 获取当前激活的工作簿
        if wb:
            return wb
        # 未检测到激活的工作簿，提示用户操作
        user = input(tip + "（打开并激活后按回车；输入 q 退出）：").strip().lower()
        if user == "q":
            print("👋 已退出。")
            sys.exit(1)


# ------------------- Word 解析辅助 -------------------
# Word中需要提取的关键标签列表（用于识别表格/段落中的关键信息）
label_tokens = [
    '超声波探伤报告','编号','试验编号','委托编号','工程名称及','施工部位','委托单位','施工单位','监理单位',
    '构件名称','检测部位','材质','板厚','仪器型号','试块','耦合剂','表面补偿','表面状况','执行处理',
    '探头型号','探伤日期','批准','审核','试验','检测单位','报告日期','检测单位名称'
]


def is_label(text: str) -> bool:
    """判断文本是否为标签（含标签关键词或为空）"""
    t = (text or '').strip()
    return not t or any(tok in t for tok in label_tokens)


def first_date_str(date_range: str) -> str:
    """从日期范围字符串中提取第一个日期（如“2023年3月1日-3月5日”→“2023年3月1日”）"""
    match: Optional[ReMatch[str]] = re.search(r"(\d{4}年\d{1,2}月\d{1,2}日)", date_range or "")
    return match.group(1) if match is not None else (date_range or '').strip()


def parse_cn_date(s: str):
    """
    解析中文日期字符串（如“2023年3月1日”）为datetime对象
    :param s: 中文日期字符串
    :return: 对应的datetime对象，解析失败返回None
    """
    match: Optional[ReMatch[str]] = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", s or "")
    if match is None:
        return None
    y, m, d = map(int, match.groups())
    try:
        return datetime(y, m, d)
    except ValueError:  # 日期无效（如2月30日）
        return None


def parse_cn_date_range(text: str):
    """
    解析中文日期范围（如“2023年3月1日至3月5日”）为开始和结束datetime对象
    :param text: 包含日期范围的文本
    :return: (开始日期, 结束日期)，解析失败返回(None, None)
    """
    parts = re.findall(r"(\d{4})年(\d{1,2})月(\d{1,2})日", text or "")  # 提取所有日期
    if not parts:
        return None, None
    if len(parts) == 1:  # 只有一个日期
        y, m, d = map(int, parts[0])
        return datetime(y, m, d), None
    # 取前两个日期作为开始和结束
    (y1, m1, d1), (y2, m2, d2) = parts[:2]
    return datetime(int(y1), int(m1), int(d1)), datetime(int(y2), int(m2), int(d2))


def value_after_label(row_cells, label_sub: str):
    """
    从表格行中提取标签后的第一个有效值（非标签文本）
    :param row_cells: 表格行的单元格文本列表
    :param label_sub: 标签关键词（如“委托编号”）
    :return: 标签后的有效值，未找到返回None
    """
    # 找到包含标签关键词的单元格索引
    indices = [i for i, t in enumerate(row_cells) if label_sub in (t or "")]
    if not indices:
        return None
    j = indices[0]
    # 从标签后一个单元格开始找第一个非标签文本
    for k in range(j + 1, len(row_cells)):
        t = (row_cells[k] or "").strip()
        if t and not is_label(t):
            return t
    return None


def extract_data_from_word(word_path: Union[str, Path]) -> dict:
    """
    从Word文档中提取探伤报告关键信息（优先表格，段落兜底）
    :param word_path: Word文件路径
    :return: 提取的关键信息字典（如委托编号、工程名称等）
    """
    doc = Document(str(word_path))
    # 提取所有段落文本（过滤空行）
    paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_texts: List[str] = []  # 存储所有表格文本（用于兜底解析）
    result: dict = {}  # 存储提取的结果

    # 优先从表格中提取信息
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells]  # 行内所有单元格文本
            table_texts.extend([c.strip() for c in cells if c.strip()])  # 收集表格文本
            # 提取目标字段（标签与目标字段映射）
            for target, label in [
                ('委托编号','委托编号'),
                ('工程名称','工程名称及'),
                ('检测部位','检测部位'),
                ('材质','材质'),
                ('探头型号','探头型号'),
                ('探伤日期','探伤日期'),
                ('执行处理','执行处理'),
            ]:
                if target not in result:  # 已提取的字段不再重复处理
                    v = value_after_label(cells, label)
                    if v:
                        result[target] = v

    # 合并表格和段落文本，用于兜底解析（表格中未提取到的信息）
    full_text = "\n".join(table_texts + paragraph_texts)

    # 提取质量等级（从全文中匹配）
    match_grade: Optional[ReMatch[str]] = re.search(r"本次检测共测试([一二])级焊缝", full_text)
    if match_grade is not None:
        result["质量等级"] = f"{match_grade.group(1)}级"
    else:
        # 兜底：直接匹配“一级”或“二级”关键词
        if "一级" in full_text:
            result["质量等级"] = "一级"
        elif "二级" in full_text:
            result["质量等级"] = "二级"

    # 探伤日期兜底（表格中未提取到时，从全文匹配）
    if not result.get("探伤日期"):
        match_date: Optional[ReMatch[str]] = re.search(r"探伤日期[:：]?\s*([0-9年月日～\-\s]+)", full_text)
        if match_date is not None:
            result["探伤日期"] = match_date.group(1).strip()
    # 格式化探伤日期（取第一个日期）
    if result.get("探伤日期"):
        result["探伤日期"] = first_date_str(result["探伤日期"])

    # 执行处理（检测依据）兜底
    if not result.get("执行处理"):
        match_proc: Optional[ReMatch[str]] = re.search(r"执行处理[:：]?\s*([A-Za-z0-9/—、，,;；\s-]+)", full_text)
        if match_proc is not None:
            result["执行处理"] = match_proc.group(1).strip()

    # 保存原始探伤日期相关文本（用于后续日期解析）
    result["_原始探伤日期串"] = full_text.split("探伤日期")[-1] if "探伤日期" in full_text else result.get("探伤日期", "")
    return result


# ------------------- 温度估计（纯数字） -------------------
# 北京月均温度（用于估计检测环境温度）
month_mean = {1:-3,2:0,3:6,4:14,5:20,6:24,7:26,8:25,9:20,10:13,11:5,12:-1}


def beijing_temp_guess_number(dt: datetime) -> str:
    """
    根据日期估计北京的检测环境温度（月均温插值+随机微扰）
    :param dt: 日期
    :return: 估计的温度字符串（整数）
    """
    if not dt:
        return ""
    month, year = dt.month, dt.year
    # 计算当月天数和下个月（用于插值）
    if month == 12:
        days_in_month = (datetime(year+1,1,1) - datetime(year,12,1)).days  # 12月天数
        next_month = 1
    else:
        days_in_month = (datetime(year,month+1,1) - datetime(year,month,1)).days  # 当月天数
        next_month = month + 1
    # 计算当月内的插值比例（0-1）
    frac = (dt.day - 1) / max(days_in_month - 1, 1)
    # 月均温插值（线性）
    base = month_mean[month] + frac * (month_mean[next_month] - month_mean[month])
    # 随机微扰（基于日期的伪随机，范围-2~2）
    wiggle = ((dt.day * 37) % 5) - 2
    return str(int(round(base + wiggle)))


def format_cn_date(dt: datetime) -> str:
    """将datetime对象格式化为中文日期字符串（如“2023年3月1日”）"""
    return f"{dt.year}年{dt.month}月{dt.day}日"


# ------------------- Excel 辅助 -------------------
def get_sheet_by_names(workbook, name_candidates: List[str]):
    """
    根据候选名称从工作簿中查找工作表（忽略名称顺序）
    :param workbook: Excel工作簿对象
    :param name_candidates: 工作表名称候选列表
    :return: 找到的工作表对象，未找到返回None
    """
    for sheet in workbook.Sheets:
        if sheet.Name in name_candidates:
            return sheet
    return None


def put_value(sheet, rng: str, value: str):
    """
    向Excel单元格或单元格区域写入值（兼容合并单元格）
    :param sheet: 工作表对象
    :param rng: 单元格区域（如“A1:B2”）
    :param value: 要写入的值
    """
    rng_obj = sheet.Range(rng)
    cell11 = rng_obj.Cells(1, 1)  # 取区域第一个单元格
    try:
        if cell11.MergeCells:  # 若为合并单元格，写入合并区域
            cell11.MergeArea.Value = value
        else:  # 非合并单元格，直接写入区域
            rng_obj.Value = value
    except Exception:  # 异常时直接写入第一个单元格
        cell11.Value = value


# ---- 检测依据分发 ----
def normalize_code(s: str) -> str:
    """标准化检测依据代码（统一大小写、替换特殊符号）"""
    return (s or "").upper().replace(" ", "").replace("—", "-").replace("－", "-")


def write_detection_basis(sheet, basis_str: str):
    """
    将检测依据（执行处理）写入Excel对应单元格（按标准代码映射）
    :param sheet: 工作表对象
    :param basis_str: 检测依据字符串（如“GB50205-2020, GB50661-2011”）
    """
    # 标准代码与Excel单元格映射
    address_map = {
        "GB50205-2020":   "B12:B12",
        "GB50661-2011":   "C12:C12",
        "JG/T203-2007":   "D12:D12",
        "GB/T50621-2010": "E12:E12",
        "GB/T11345-2023": "F12:F12",
        "GB/T29712-2023": "G12:G12",
        "GB/T29711-2023": "H12:H12",
    }
    # 先清空目标区域
    for addr in ["B12:B12","C12:C12","D12:D12","E12:E12","F12:F12","G12:G12","H12:H12","I12:J12"]:
        put_value(sheet, addr, "")
    if not basis_str:
        return
    # 拆分检测依据为多个 token
    tokens = [t for t in re.split(r"[、，,;；\s]+", basis_str) if t.strip()]
    unknown_list: List[str] = []  # 存储未识别的代码
    used: set = set()  # 记录已使用的标准代码（去重）
    for token in tokens:
        norm = normalize_code(token).replace("GBT","GB/T")  # 标准化并修正GB/T格式
        if norm in address_map and norm not in used:
            put_value(sheet, address_map[norm], token.strip())
            used.add(norm)
        else:
            unknown_list.append(token.strip())
    # 未识别的代码写入备用单元格
    if unknown_list:
        put_value(sheet, "I12:J12", ", ".join(unknown_list))


# ------------------- 探头自动选型（按 D列板厚 + C列焊缝类型） -------------------
# 探头选型规则：键为焊缝类型（'D'对接；'JD'角对接），值为(板厚下限, 板厚上限, 探头清单)
PROBE_RULES: dict[str, List[Tuple[float, float, List[str]]]] = {
    "D": [
        (8, 15,  ["A2.5P9×9A70°"]),
        (15, 25, ["A2.5P9×9A70°"]),
        (25, 40, ["A2.5P9×9A70°", "A2.5P9×9A45°"]),
        (40, 50, ["A2.5P9×9A60°", "A2.5P9×9A45°"]),
        (50, 75, ["A2.5P13×13A70°", "A2.5P13×13A45°"]),
        (75, 100,["A2.5P13×13A60°", "A2.5P13×13A45°"]),
        (100, 1e9,["A2.5P13×13A60°", "A2.5P13×13A45°"]),  # 1e9表示≥100
    ],
    "JD": [
        (8, 15,  ["A2.5P9×9A70°"]),
        (15, 25, ["A2.5P9×9A70°"]),
        (25, 40, ["A2.5P9×9A60°", "A2.5P9×9A45°"]),
        (40, 50, ["A2.5P9×9A70°", "A2.5P9×9A60°"]),
        (50, 75, ["A2.5P13×13A70°", "A2.5P13×13A60°", "A2.5P13×13A45°"]),
        (75, 100,["A2.5P13×13A70°", "A2.5P13×13A60°", "A2.5P13×13A45°"]),
        (100, 1e9,["A2.5P9×9A70°", "A2.5P13×13A70°", "A2.5P13×13A60°", "A2.5P13×13A45°"]),
    ],
}


def _to_float(x) -> Optional[float]:
    """将输入转换为浮点数（提取数字部分），失败返回None"""
    if x is None:
        return None
    s = str(x).strip()
    m = re.search(r"(-?\d+(?:\.\d+)?)", s)  # 匹配数字（整数或小数）
    if not m:
        return None
    try:
        return float(m.group(1))
    except ValueError:
        return None


def _pick_weld_type(cell_c, cell_alt=None) -> Optional[str]:
    """
    根据单元格内容判断焊缝类型（优先C列，兼容备用列）
    :param cell_c: C列单元格值
    :param cell_alt: 备用列单元格值
    :return: 焊缝类型（'D'对接；'JD'角对接），未识别返回None
    """
    for c in (cell_c, cell_alt):
        t = str(c or "").upper()
        if "JD" in t:
            return "JD"
    for c in (cell_c, cell_alt):
        t = str(c or "").upper()
        if "D" in t:
            return "D"
    return None


def choose_probes_for_segments(data_excel_path: Path,
                               segments: Optional[List[Tuple[int, int]]]) -> List[str]: # noqa
    """
    根据数据Excel中的板厚和焊缝类型自动选择探头型号
    :param data_excel_path: 数据Excel路径
    :param segments: 行段列表（(开始行, 结束行)），None表示全部行
    :return: 选好的探头型号列表（去重，最多8个）
    """
    if not data_excel_path.exists():
        return []

    excel_app = win32.Dispatch("Excel.Application")
    wb = excel_app.Workbooks.Open(str(data_excel_path), ReadOnly=True)  # 只读打开
    xl_up = -4162  # Excel常量：向上查找
    try:
        ws = wb.Sheets(1)  # 取第一个工作表
        last_row = ws.Cells(ws.Rows.Count, 2).End(xl_up).Row  # B列最后一行有数据的行号
        if segments is None:
            segments = [(2, last_row)]  # 默认为第2行到最后一行

        picked: List[str] = []  # 选中的探头型号
        seen = set()  # 去重集合

        # 遍历所有行段
        for (s, e) in segments:
            s = max(s, 2); e = min(e, last_row)  # 确保行号有效
            for r in range(s, e + 1):
                # 获取焊缝类型（C列）
                weld_type = _pick_weld_type(ws.Cells(r, 3).Value)
                if weld_type not in PROBE_RULES:
                    continue
                # 获取板厚（D列优先，E列兜底）
                thickness = _to_float(ws.Cells(r, 4).Value) or _to_float(ws.Cells(r, 5).Value)
                if thickness is None:
                    continue
                # 根据板厚匹配探头
                for low, high, models in PROBE_RULES[weld_type]:
                    if (thickness >= low) and (thickness < high):
                        for m in models:
                            if m not in seen:
                                picked.append(m); seen.add(m)
                        break  # 匹配到后跳出循环
        return picked[:8]  # 最多返回8个
    finally:
        wb.Close(SaveChanges=False)  # 关闭工作簿，不保存


# ------------------- 写第一张：基础信息 -------------------
def fill_sheet_basic(workbook, data: dict,
                     override_date: Optional[datetime] = None,
                     override_probes: Optional[List[str]] = None):
    """
    填充Excel中的“基础信息”工作表
    :param workbook: 目标工作簿对象
    :param data: 从Word提取的信息字典
    :param override_date: 覆盖的日期（多日期场景下使用）
    :param override_probes: 覆盖的探头列表（自动选型结果）
    """
    # 获取“基础信息”工作表（兼容“基础性息”错别字）
    sheet = get_sheet_by_names(workbook, ["基础性息", "基础信息"]) or workbook.Sheets(1)

    # 基础信息单元格映射（字段→Excel区域）
    cols = {
        "工程名称":"B2:D2","委托编号":"F2:J2","仪器编号":"B3:D3","检测环境":"B4:D4","探伤日期":"F3:J3",
        "测试人":"B5:D5","K值角度":"F5:J5","材质":"B6:D6","表面补偿":"F7:J7","检测灵敏度":"F8:J8",
        "表面状况":"B9:D9","质量等级":"F9:G9","焊接方式":"B10:D10","检测时机":"F10:J10","坡口形式":"F11:J11",
        "备注":"B21:J21",
    }

    # 处理日期相关逻辑
    date_from_word = parse_cn_date(data.get("探伤日期") or "")  # 从Word提取的日期
    chosen_date = override_date or date_from_word  # 最终使用的日期（多日期场景下用override_date）
    # 仪器编号判断（根据日期范围）
    window_start = datetime(2025, 3, 12)
    window_end = datetime(2025, 4, 9)
    instrument = "13-27" if chosen_date and (window_start <= chosen_date <= window_end) else "22-72"
    # 估计检测环境温度
    temp_number = beijing_temp_guess_number(chosen_date)

    # 填充固定值字段
    put_value(sheet, cols["仪器编号"], instrument)
    if temp_number: put_value(sheet, cols["检测环境"], temp_number)
    put_value(sheet, cols["测试人"], "于征")
    put_value(sheet, cols["K值角度"], "角度")
    put_value(sheet, cols["表面补偿"], "4dB")
    put_value(sheet, cols["检测灵敏度"], "DAC-14dB")
    put_value(sheet, cols["表面状况"], "磨光")
    put_value(sheet, cols["焊接方式"], "气保")
    put_value(sheet, cols["检测时机"], "焊后24h")
    put_value(sheet, cols["坡口形式"], "L")

    # 填充从Word提取的字段
    if v := data.get("工程名称"): put_value(sheet, cols["工程名称"], v)
    if v := data.get("委托编号"): put_value(sheet, cols["委托编号"], v)
    if v := data.get("材质"):    put_value(sheet, cols["材质"], v)
    if chosen_date:              put_value(sheet, cols["探伤日期"], format_cn_date(chosen_date))
    if v := data.get("质量等级"): put_value(sheet, cols["质量等级"], v)

    # 处理扫查方式和备注（根据检测部位判断）
    part = data.get("检测部位", "")
    has_corner = "角对接焊缝" in part
    has_butt = "对接焊缝" in part
    if has_corner and has_butt:
        sheet.Range("B11").Value = "单面单侧"
        sheet.Range("C11").Value = "单面双侧"
        put_value(sheet, cols["备注"], "注：D表示对接、JD表示角对接")
    elif has_corner:
        sheet.Range("B11").Value = "单面单侧"
        put_value(sheet, cols["备注"], "注：JD表示角对接")
    else:
        sheet.Range("B11").Value = "单面双侧"
        put_value(sheet, cols["备注"], "注：D表示对接")

    # 填充试块信息
    sheet.Range("F6").Value = "CSK-IA"
    sheet.Range("H6").Value = "RB-1"
    sheet.Range("I6").Value = "RB-2"

    # 填充耦合剂
    put_value(sheet, "B7:D7", "化学浆糊")
    put_value(sheet, "B8:D8", "化学浆糊")

    # 填充检测依据
    write_detection_basis(sheet, data.get("执行处理", ""))

    # 填充探头型号（优先自动选型，其次Word提取，最多8个）
    for i in range(13, 21):  # 清空B13到B20
        sheet.Range(f"B{i}").Value = ""
    if override_probes is not None:
        for i, p in enumerate(override_probes[:8]):
            sheet.Range(f"B{13+i}").Value = p
    else:
        # 从Word提取的探头型号拆分后填充
        probes = re.split(r"[、,，;；\s]+", (data.get("探头型号") or "").strip())
        probes = [p for p in probes if p]
        for i, p in enumerate(probes[:8]):
            sheet.Range(f"B{13+i}").Value = p


# ------------------- 第二张：数据信息（按需插行 + 进度条） -------------------
def fill_sheet_data(workbook_server,
                    data_excel_path: Path,
                    row_ranges: Optional[List[Tuple[int, int]]] = None) -> None:
    """
    填充Excel中的“数据信息”工作表（从外部数据Excel拷贝数据）
    :param workbook_server: 目标工作簿对象
    :param data_excel_path: 外部数据Excel路径
    :param row_ranges: 要拷贝的行段列表，None表示全部行
    """
    # 定义Excel常量（兼容不同版本）
    try:
        const = win32.constants
        xl_up = const.xlUp  # 向上查找
        xl_paste_formats = const.xlPasteFormats  # 粘贴格式
        color_index_none = getattr(const, 'xlColorIndexNone', -4142)  # 无颜色
        pattern_none = getattr(const, 'xlPatternNone', -4142)  # 无填充
    except Exception:
        xl_up = -4162
        xl_paste_formats = -4122
        color_index_none = -4142
        pattern_none = -4142

    # 获取“数据信息”工作表
    ws_dst = get_sheet_by_names(workbook_server, ["数据信息"]) or workbook_server.Sheets(2)
    app_dst = workbook_server.Application  # 目标Excel应用实例
    ws_dst.Activate()  # 激活工作表

    # 检查数据文件是否存在
    if not data_excel_path.exists():
        print_step(f"⚠️ 没找到外部数据文件：{data_excel_path}，已跳过『数据信息』拷贝。")
        return

    # 打开外部数据Excel（只读）
    app_src = win32.Dispatch("Excel.Application")
    wb_src = app_src.Workbooks.Open(str(data_excel_path), ReadOnly=True)

    try:
        ws_src = wb_src.Sheets(1)  # 取第一个工作表
        # 获取源表B列最后一行有数据的行号
        last_row_src = ws_src.Cells(ws_src.Rows.Count, 2).End(xl_up).Row
        if row_ranges is None:
            if last_row_src < 2:  # 无数据
                print_step("⚠️ 源表 B 列从第2行起没有数据，跳过复制。")
                return
            row_ranges = [(2, last_row_src)]  # 默认为第2行到最后一行

        # 自动探测“备注”列（查找表头为“备注”的列）
        remark_col = None
        for probe_row in (1, 2):  # 检查第1行和第2行
            for c in range(1, 30):  # 前30列
                title = str(ws_src.Cells(probe_row, c).Value or "").strip()
                if title == "备注":
                    remark_col = c
                    break
            if remark_col:
                break

        # 收集要拷贝的数据（A、D、E、F列对应源表B、C、D、E列，M列对应备注列）
        rows: List[Tuple[object, object, object, object, object]] = []
        total_rows = 0
        for s, e in row_ranges:
            if e < s:
                continue
            for r in range(s, e + 1):
                a_val = ws_src.Cells(r, 2).Value   # 目标A列 → 源表B列
                d_val = ws_src.Cells(r, 3).Value   # 目标D列 → 源表C列
                e_val = ws_src.Cells(r, 4).Value   # 目标E列 → 源表D列（板厚）
                f_val = ws_src.Cells(r, 5).Value   # 目标F列 → 源表E列（角度/回波）
                m_val = ws_src.Cells(r, remark_col).Value if remark_col else None  # 目标M列 → 源表备注列
                rows.append((a_val, d_val, e_val, f_val, m_val))
            total_rows += (e - s + 1)

        if total_rows <= 0:
            print_step("⚠️ 没有需要拷贝的行。")
            return

        # 静默模式写入数据（关闭弹窗和刷新）
        with excel_quiet(app_dst):
            print_step(f"准备写入 {total_rows} 行数据…")
            progress_bar(0, total_rows, prefix="写入")

            start_row = 3  # 从第3行开始写入
            end_row = start_row + total_rows - 1  # 计算结束行

            # 清空目标区域旧内容（保留格式）
            clear_to = max(end_row, 200 + 2)  # 至少清到202行
            ws_dst.Range(f"A{start_row}:M{clear_to}").ClearContents()

            # 复制第3行格式到目标行（确保格式一致）
            try:
                ws_dst.Range("A3:M3").Copy()
                ws_dst.Range(f"A3:M{end_row}").PasteSpecial(Paste=xl_paste_formats)
                app_dst.CutCopyMode = False  # 清除剪贴板
            except Exception:
                # 格式复制失败时，尝试清除填充色（兜底）
                try:
                    rng = ws_dst.Range(f"A3:M{end_row}")
                    try:
                        rng.Interior.Pattern = pattern_none
                    except Exception:
                        rng.Interior.ColorIndex = color_index_none
                except Exception:
                    pass

            # 批量写入数据（提高效率）
            ws_dst.Range(f"A{start_row}:A{end_row}").Value = [[r[0]] for r in rows]  # A列
            ws_dst.Range(f"D{start_row}:D{end_row}").Value = [[r[1]] for r in rows]  # D列
            ws_dst.Range(f"E{start_row}:E{end_row}").Value = [[r[2]] for r in rows]  # E列
            ws_dst.Range(f"F{start_row}:F{end_row}").Value = [[r[3]] for r in rows]  # F列
            if remark_col:  # 写入备注列
                ws_dst.Range(f"M{start_row}:M{end_row}").Value = [[r[4]] for r in rows]

            # L列标注“Ⅰ”（有数据的行）
            done = 0
            for i in range(total_rows):
                a_val, d_val, e_val, f_val, _ = rows[i]
                # 只要有一个字段有值，就标注“Ⅰ”
                if any(v not in (None, "", " ") for v in (a_val, d_val, e_val, f_val)):
                    ws_dst.Cells(start_row + i, 12).Value = "Ⅰ"  # L列是第12列
                done += 1
                # 更新进度条
                if (done % 100 == 0) or (done == total_rows):
                    progress_bar(done, total_rows, prefix="写入")

            # 扩展命名区域（方便后续打印/处理）
            try:
                name_candidates = {"数据区", "明细区", "DataRange", "DataArea", "数据范围"}
                nm = None
                for n in list(workbook_server.Names):  # 遍历工作簿中的命名区域
                    n_name = getattr(n, "NameLocal", None) or getattr(n, "Name", "")
                    if any(cand in n_name for cand in name_candidates):
                        nm = n; break
                if nm is not None:
                    ws_name = ws_dst.Name
                    nm.RefersTo = f"={ws_name}!$A$3:$M${end_row}"  # 更新区域范围
            except Exception:
                pass

        print_step("『数据信息』已填完（静默写入，保存/打印不弹窗，备注也同步）。")

    finally:
        wb_src.Close(SaveChanges=False)  # 关闭源数据工作簿，不保存


# ------------------- 从“数据.xlsx”识别多日期行段（按 F 列日期标记） -------------------
def parse_day_ranges_from_source(data_excel_path: Path,
                                 year_hint: Optional[int]) -> List[Tuple[datetime, List[Tuple[int, int]]]]:     # noqa
    """
    从数据Excel的F列识别日期标记，拆分不同日期对应的行段
    :param data_excel_path: 数据Excel路径
    :param year_hint: 年份提示（用于补全日期）
    :return: 按日期排序的列表，每个元素为(日期, 行段列表)
    """
    if not data_excel_path.exists():
        return []
    excel_app = win32.Dispatch("Excel.Application")
    workbook = excel_app.Workbooks.Open(str(data_excel_path), ReadOnly=True)
    go_up = -4162  # 向上查找
    try:
        sheet = workbook.Sheets(1)
        # 获取F列最后一行有数据的行号
        last_row = sheet.Cells(sheet.Rows.Count, 6).End(go_up).Row
        if last_row < 2:
            return []
        # 匹配日期格式（如“3.31”“4/4”“4月4日”）
        pattern = re.compile(r"(\d{1,2})\s*[./月]\s*(\d{1,2})")
        markers: List[Tuple[int, int, int]] = []  # (行号, 月, 日)
        for r in range(2, last_row + 1):
            val = sheet.Cells(r, 6).Value  # F列值
            if val is None:
                continue
            txt = str(val).strip().replace("．", ".").replace("。", ".")  # 替换全角符号
            m: Optional[ReMatch[str]] = pattern.search(txt)
            if m is not None:
                month, day = int(m.group(1)), int(m.group(2))
                markers.append((r, month, day))
        if not markers:  # 未找到日期标记
            return []

        # 按日期标记拆分行段
        prev_row = 1
        ranges_per_day: List[Tuple[Tuple[int, int], Tuple[int, int]]] = []
        for row, month, day in markers:
            ranges_per_day.append(((month, day), (prev_row + 1, row)))
            prev_row = row
        # 按日期分组行段
        day_to_ranges: dict = {}
        for (month, day), (start_row, end_row) in ranges_per_day:
            day_to_ranges.setdefault((month, day), []).append((start_row, end_row))

        # 补全年份并转换为datetime
        if year_hint is None:
            year_hint = datetime.now().year
        out: List[Tuple[datetime, List[Tuple[int, int]]]] = []
        for (month, day), ranges in day_to_ranges.items():
            try:
                dt = datetime(year_hint, month, day)
                out.append((dt, ranges))
            except ValueError:  # 无效日期（如2月30日）
                continue
        out.sort(key=lambda x: x[0])  # 按日期排序
        return out
    finally:
        workbook.Close(SaveChanges=False)


# ------------------- 主入口 -------------------
def main():
    """程序主入口：引导用户输入路径，解析数据，填充Excel模板"""
    # 1) 获取Word文件路径（默认路径+用户输入）
    default_word = Path(r"D:\eg.docx")
    word_file = prompt_path_with_retry("📄 Word 路径", default_word, must_exist=True, allowed_suffixes={".docx"})

    # 2) 获取外部数据Excel路径（默认桌面+Word文件名+“数据”）
    desktop = get_desktop_path()
    default_data_excel = desktop / f"{word_file.stem}数据.xlsx"
    data_excel_path = prompt_path_with_retry("📊 外部“数据”Excel 路径",
                                             default_data_excel,
                                             must_exist=True,
                                             allowed_suffixes={".xlsx", ".xls"})

    # 3) 解析Word文档，提取关键信息
    data_from_word = extract_data_from_word(word_file)
    # 获取年份提示（从Word提取的日期中获取）
    start_date, _ = parse_cn_date_range(data_from_word.get("_原始探伤日期串", ""))
    base_year = (start_date or parse_cn_date(data_from_word.get("探伤日期") or "") or datetime.now()).year

    # 4) 从数据Excel中识别多日期行段
    day_groups = parse_day_ranges_from_source(data_excel_path, base_year)

    # 5) 连接Excel应用，处理单日期/多日期场景
    excel_app = win32.Dispatch("Excel.Application")

    # 单日期场景
    if not day_groups:
        workbook_server = wait_active_workbook(excel_app, "❗没检测到活动工作簿，请先在 Excel 服务器里打开模板并激活")
        # 自动选择探头
        probes_today = choose_probes_for_segments(data_excel_path, None)
        # 填充基础信息和数据信息
        fill_sheet_basic(workbook_server, data_from_word,
                         override_date=parse_cn_date(data_from_word.get("探伤日期") or ""),
                         override_probes=probes_today)
        fill_sheet_data(workbook_server, data_excel_path, row_ranges=None)
        print("✅ 完成：单日期填报。回到 Excel 点『打印』即可。")
        return

    # 多日期场景：逐天处理
    total_days = len(day_groups)
    print_step(f"总共 {total_days} 天，按日期升序逐天填报。")
    progress_bar(0, total_days, prefix="整体进度")

    for idx, (dt, ranges) in enumerate(day_groups, 1):
        if idx == 1:  # 第一天：初始化
            print(f"\n▶ 开始填写第 {idx}/{total_days} 天：{format_cn_date(dt)}")
            workbook_server = wait_active_workbook(excel_app, "❗请先在 Excel 服务器里打开模板并激活")
        else:  # 后续天数：提示用户新建模板
            user = input(f"\n⏸ 第 {idx-1} 天已填完。请『新建下一张模板』并激活，然后按回车继续（输入 q 退出）：").strip().lower()
            if user == "q":
                print("👋 已退出。"); sys.exit(0)
            workbook_server = wait_active_workbook(excel_app, "❗仍未检测到活动工作簿，请激活模板")

        # 按当天行段选择探头
        probes_today = choose_probes_for_segments(data_excel_path, ranges)
        # 填充当天的基础信息和数据信息
        fill_sheet_basic(workbook_server, data_from_word,
                         override_date=dt,
                         override_probes=probes_today)
        fill_sheet_data(workbook_server, data_excel_path, row_ranges=ranges)

        print(f"✅ 已填写：{format_cn_date(dt)}")
        progress_bar(idx, total_days, prefix="整体进度")
        # 提示下一天操作
        if idx < total_days:
            next_dt = day_groups[idx][0]
            print(f"👉 下一天：{format_cn_date(next_dt)}（先在 Excel 里创建下一张，再回来按回车）")

    print("\n🎉 全部日期都填完了。回到 Excel 按『打印』出报告吧。")


if __name__ == "__main__":
    main()

                                                                                                        # v 4.1.1